# الأعلام الجغرافية الفلسطينية بين الطمس والتحريف
# وثائق من الأرشيف الصهيوني

import urllib.request
import os
import json
import docx
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.backends.backend_pdf
from matplotlib.patches import ConnectionPatch

base_url = "https://sys.archives.gov.il/api/solr?fq[]=lang_code_s:he&fq[]=product_id_i:"
base_attachment_url = "https://www.archives.gov.il/"

attachment_dir = "./Archives/Attachment/"
data_analysis_dir = "./Archives/DataAnalysis/"
metadata_dir = "./Archives/MetaData/"

start_year_to_filter = 1948

status_count = [0, 0, 0]
d_type = {}
documentetionSource = {}
d_year = {'[1948, 1950)': 0, '[1950, 1960)': 0, '[1960, 1970)': 0, '[1970, 1980)': 0,
          '[1980, 1990)': 0, '[1990, 2000)': 0, '[2000, 2010)': 0, '[2010, 2020)': 0, '[2020, ...)': 0}

# init docx
header = ['סוג הקובץ', 'סטטוס חשיפה', 'תקופת החומר עד', 'תקופת החומר מ', 'שם הקובץ']
doc = docx.Document()
doc.add_heading('רשימת הקבצים שנמצאו',0)
docTable = doc.add_table(rows=1,cols=5)
docTable.style = 'Medium Shading 2 Accent 3'
hdr_Cells = docTable.rows[0].cells
for i in range(0, 5):
    hdr_Cells[i].text = header[i]
###########

def read_jsons(start_at, do_it_for):
    print("gggggggggg") #
    b = 0
    for i in range(start_at, start_at + do_it_for):
        b = b + 1
        if(b % 100 == 0):
            print(b)
        #
        url = base_url + str(i)
        try:
            initial_product_meta_data = read_meta_data(url)
            start_year = int(initial_product_meta_data[0]["objDate_datingPeriodStartYear_t"])
            addAttr = str(initial_product_meta_data[0]["addAttr_orgTree_t"])
            addAttr_ = str(initial_product_meta_data[0]["addAttr_objectName_s"])
            file_name = initial_product_meta_data[0]["objDesc_objectName_t"]
            desc = str(initial_product_meta_data[0]["objDesc_objectDesc_t"])
            
            if((start_year > start_year_to_filter) & ("ועדת השמות הממשלתית" in [addAttr, addAttr_, file_name, desc])):
                datingPeriodStart = str(initial_product_meta_data[0]["objDate_datingPeriodStart_t"])
                datingPeriodEnd = str(initial_product_meta_data[0]["objDate_datingPeriodEnd_t"])
                status = str(initial_product_meta_data[0]["addAttr_statusChasifa_t"])
                attachmentType = str(initial_product_meta_data[0]["objHier_attachment_attachmentType_s"])

                add_row_to_table(attachmentType, status, datingPeriodEnd, datingPeriodStart, file_name)
                update_dic_year(int(start_year/10)*10)
                update_dic(documentetionSource, str(initial_product_meta_data[0]["objHier_archiveName_t"]) [::-1])
                load_data_into_file(initial_product_meta_data, file_name + ".json")

                if("גלוי" in status):
                    load_attachments(initial_product_meta_data[0], file_name)
                    update_dic(d_type, attachmentType)
                    status_count[0] += 1
                elif("חסוי" in status):
                    status_count[1] += 1
                else: # נחשף בחלקו
                    status_count[2] += 1
        except:
            continue

def add_row_to_table(attachmentType, status, datingPeriodEnd, datingPeriodStart, file_name):
    row_Cells = docTable.add_row().cells
    row_Cells[0].text = attachmentType
    row_Cells[1].text = status
    row_Cells[2].text = datingPeriodEnd
    row_Cells[3].text = datingPeriodStart
    row_Cells[4].text = file_name

def update_dic_year(round_yead):
    if round_yead < 1950:
        d_key = '[1948, 1950)'
    elif round_yead > 2010:
        d_key = '[2020, ...)'
    else:
        d_key = '[' + str(round_yead) + ', ' + str(round_yead+10) + ')'
    d_year[d_key] += 1

def update_dic(dic, key):
    if key in dic:
        dic[key] += 1
    else:
        dic[key] = 1

# Generic function to read json data into an object from a give url
def read_meta_data(url):
    try:
        data = urllib.request.urlopen(url)
        if is_valid_data(data):
            metadata = json.loads(data.read())
            docs_data = metadata["response"]["docs"]
            return docs_data
    except:
        raise Exception("Could not open url")

def is_valid_data(data):
    return data.getcode() != 404

# Load the meta data as a json file into external/local location
def load_data_into_file(data_to_load, file_name):
    try:
        with open(metadata_dir + file_name, 'w', encoding='utf8') as jfile:
            jfile.write(json.dumps(data_to_load, indent=4, ensure_ascii=False))
            jfile.close()
    except:
        pass

# Load files that we've retrieved from the given url into external/local location
def load_attachments(product_data, file_name):
    try:
        file_extension = "." + product_data["objHier_attachment_attachmentType_s"]
        url = base_attachment_url + product_data["attachment_url_s"]
        urllib.request.urlretrieve(url, attachment_dir + file_name + file_extension)
    except:
        raise Exception("Could not load the wanted attachment")

def statusDA():
    # make figure and assign axis objects
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 5))
    fig.subplots_adjust(wspace=0)

    # pie chart parameters
    ratios = status_count
    labels = ['Accessible', 'Hidden', 'Partially exposed']
    explode = [0.1, 0, 0] # only "explode" the 1st slice (i.e. 'Accessible')
    # rotate so that first wedge is split by the x-axis
    angle = ratios[0] # -180 * ratios[0]
    ax1.pie(ratios, autopct='%1.1f%%', startangle=angle, labels=labels,
            explode=explode, colors=['lightskyblue', 'lightcoral', 'gold'])
    ax1.set_title(str(ratios[0]+ratios[1]+ratios[2])+' files downloaded')

    # bar chart parameters
    xpos = 0
    bottom = 0
    ratios = list(d_type.values())
    width = .2
    colors = ['lightslategray', 'steelblue', 'dodgerblue', 'deepskyblue']

    for j in range(len(ratios)):
        ratios[j] = ratios[j] / status_count[0]
        height = ratios[j]
        ax2.bar(xpos, height, width, bottom=bottom, color=colors[j % 4])
        ypos = bottom + ax2.patches[j].get_height() / 2
        bottom += height
        ax2.text(xpos, ypos, "%d%%" % (ax2.patches[j].get_height() * 100),
                ha='center')

    ax2.set_title('File Types')
    ax2.legend(tuple(d_type.keys()))
    ax2.axis('off')
    ax2.set_xlim(- 2.5 * width, 2.5 * width)

    # use ConnectionPatch to draw lines between the two plots
    # get the wedge data
    theta1, theta2 = ax1.patches[0].theta1, ax1.patches[0].theta2
    center, r = ax1.patches[0].center, ax1.patches[0].r
    bar_height = sum([item.get_height() for item in ax2.patches])

    # draw top connecting line
    x = r * np.cos(np.pi / 180 * theta2) + center[0]
    y = r * np.sin(np.pi / 180 * theta2) + center[1]
    con = ConnectionPatch(xyA=(-width / 2, bar_height), coordsA=ax2.transData,
                        xyB=(x, y), coordsB=ax1.transData)
    con.set_color([0, 0, 0])
    con.set_linewidth(4)
    ax2.add_artist(con)

    # draw bottom connecting line
    x = r * np.cos(np.pi / 180 * theta1) + center[0]
    y = r * np.sin(np.pi / 180 * theta1) + center[1]
    con = ConnectionPatch(xyA=(-width / 2, 0), coordsA=ax2.transData,
                        xyB=(x, y), coordsB=ax1.transData)
    con.set_color([0, 0, 0])
    ax2.add_artist(con)
    con.set_linewidth(4)
    fig.savefig(data_analysis_dir + 'Status&Types.png')

def yearsDA():
    # A python dictionary
    years_data = {"Numbers": list(d_year.keys()),
                  "Files": list(d_year.values())}
    # Dictionary loaded into a DataFrame
    yearsDF = pd.DataFrame(data=years_data)
    # Draw a vertical bar chart
    fig = yearsDF.plot.bar(x="Numbers", y="Files", rot=20, title="Number of files downloaded")
    plt.savefig(data_analysis_dir + 'Years.png')

def docSource():
    # Pie chart, where the slices will be ordered and plotted counter-clockwise:
    labels = list(documentetionSource.keys())
    sizes = list(documentetionSource.values())
    fig1, ax1 = plt.subplots()
    ax1.pie(sizes, labels=labels, autopct='%1.1f%%',
            shadow=True, startangle=90, colors=matplotlib.colors.TABLEAU_COLORS)
    ax1.axis('equal') # Equal aspect ratio ensures that pie is drawn as a circle.
    ax1.set_title('Documentetion source')
    plt.savefig(data_analysis_dir + 'DocSource.png')

# Create target directorys if they don't exist
def mkdirs(path, arr_names):
    for i in range(0, len(arr_names)):
        dirName = path + '/' + arr_names[i]
        if not os.path.exists(dirName):
            os.mkdir(dirName)

def main():
    mkdirs('.', ['Archives'])
    mkdirs('./Archives', ['Attachment', 'DataAnalysis', 'MetaData'])

    read_jsons(121500, 2000)
    read_jsons(141000, 2000)
    read_jsons(151000, 2000)
    read_jsons(161000, 10000)
    read_jsons(222000, 2000)
    read_jsons(492000, 2000)
    read_jsons(510000, 90000)
    read_jsons(1367000, 2000)
    read_jsons(1735000, 2000)
    read_jsons(2182500, 2000)
    read_jsons(2238000, 2000)
    read_jsons(1508000, 2000)

    # Data analysis
    doc.save(data_analysis_dir + 'Table.docx')
    statusDA()
    yearsDA()
    docSource()
    #plt.show()
    

if __name__ == '__main__':

    main()
